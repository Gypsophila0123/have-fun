#!/usr/bin/env python3
"""
generate_word.py — 计算机犯罪类案检索报告 Word 文档生成器

用途：Phase 5 将并行写作的 Markdown 内容转换为格式化 Word 文档
依赖：python-docx (优先) 或 纯标准库 zipfile (回退方案)

用法：
    python generate_word.py --input report.md --output report.docx [--config config.json]
"""

import argparse
import json
import re
import sys
import zipfile
from pathlib import Path


DEFAULT_CONFIG = {
    "title_font": "黑体",
    "title_size": 15,
    "heading1_font": "黑体",
    "heading1_size": 15,
    "heading2_font": "黑体",
    "heading2_size": 14,
    "heading3_font": "黑体",
    "heading3_size": 12,
    "heading4_font": "宋体",
    "heading4_size": 12,
    "body_font": "宋体",
    "body_size": 12,
    "table_font": "宋体",
    "table_size": 12,
    "blockquote_font": "楷体",
    "blockquote_size": 12,
    "line_spacing": 1.25,
    "margin_top": 2.54,
    "margin_bottom": 2.54,
    "margin_left": 3.17,
    "margin_right": 3.07,
    "first_line_indent": 0.74,
    "table_border_width": 1,
    "table_cell_padding": 10,
}


def parse_markdown_to_elements(md_content):
    """Parse markdown content into structured elements."""
    elements = []
    lines = md_content.split("\n")
    i = 0
    in_table = False
    table_header = []
    table_rows = []

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            if in_table and table_header:
                elements.append({
                    "type": "table",
                    "header": table_header,
                    "rows": table_rows
                })
                table_header = []
                table_rows = []
                in_table = False
            i += 1
            continue

        if line.startswith("|") and "|" in line[1:]:
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if not in_table:
                table_header = cells
                in_table = True
            elif all(c.replace("-", "").replace(":", "").strip() == "" for c in cells):
                pass
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table and table_header:
                elements.append({
                    "type": "table",
                    "header": table_header,
                    "rows": table_rows
                })
                table_header = []
                table_rows = []
                in_table = False

        if line.startswith("# "):
            elements.append({"type": "heading", "level": 0, "text": line[2:]})
        elif line.startswith("## "):
            elements.append({"type": "heading", "level": 1, "text": line[3:]})
        elif line.startswith("### "):
            elements.append({"type": "heading", "level": 2, "text": line[4:]})
        elif line.startswith("#### "):
            elements.append({"type": "heading", "level": 3, "text": line[5:]})
        elif line.startswith("- ") or line.startswith("* "):
            elements.append({"type": "list", "text": line[2:]})
        elif line.startswith("**") and line.endswith("**"):
            elements.append({"type": "bold", "text": line.strip("*")})
        elif line.startswith("> "):
            elements.append({"type": "quote", "text": line[2:]})
        else:
            elements.append({"type": "paragraph", "text": line})

        i += 1

    if in_table and table_header:
        elements.append({
            "type": "table",
            "header": table_header,
            "rows": table_rows
        })

    return elements


def escape_xml(text):
    """Escape XML special characters."""
    if not text:
        return ""
    text = str(text)
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    text = text.replace('"', "&quot;")
    text = text.replace("'", "&apos;")
    return text


def generate_ooxml_document(elements, config):
    """Generate OOXML document.xml content from elements."""
    body_parts = []

    for elem in elements:
        if elem["type"] == "heading":
            level = elem["level"]
            text = escape_xml(elem["text"])
            if level == 0:
                font_name = config.get("title_font", "黑体")
                font_size = config.get("title_size", 15) * 2
            elif level == 1:
                font_name = config.get("heading1_font", "黑体")
                font_size = config.get("heading1_size", 15) * 2
            elif level == 2:
                font_name = config.get("heading2_font", "黑体")
                font_size = config.get("heading2_size", 14) * 2
            elif level == 3:
                font_name = config.get("heading3_font", "黑体")
                font_size = config.get("heading3_size", 12) * 2
            else:
                font_name = config.get("heading4_font", "宋体")
                font_size = config.get("heading4_size", 12) * 2
            
            body_parts.append(f"""<w:p>
  <w:pPr>
    <w:pStyle w:val="Heading{level + 1}"/>
    <w:spacing w:after="200"/>
  </w:pPr>
  <w:r>
    <w:rPr>
      <w:rFonts w:eastAsia="{font_name}" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
      <w:b/>
      <w:sz w:val="{font_size}"/>
      <w:szCs w:val="{font_size}"/>
    </w:rPr>
    <w:t>{text}</w:t>
  </w:r>
</w:p>""")

        elif elem["type"] == "paragraph":
            text = escape_xml(elem["text"])
            font_name = config.get("body_font", "宋体")
            font_size = config.get("body_size", 12) * 2
            line_spacing = int(config.get("line_spacing", 1.25) * 240)
            first_indent = int(config.get("first_line_indent", 0.74) * 28.35 * 100)
            
            body_parts.append(f"""<w:p>
  <w:pPr>
    <w:spacing w:line="{line_spacing}" w:lineRule="auto"/>
    <w:jc w:val="both"/>
    <w:ind w:firstLine="{first_indent}"/>
  </w:pPr>
  <w:r>
    <w:rPr>
      <w:rFonts w:eastAsia="{font_name}" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
      <w:sz w:val="{font_size}"/>
      <w:szCs w:val="{font_size}"/>
    </w:rPr>
    <w:t>{text}</w:t>
  </w:r>
</w:p>""")

        elif elem["type"] == "list":
            text = escape_xml(elem["text"])
            font_name = config.get("body_font", "宋体")
            font_size = config.get("body_size", 12) * 2
            
            body_parts.append(f"""<w:p>
  <w:pPr>
    <w:spacing w:line="240" w:lineRule="auto"/>
    <w:numPr>
      <w:ilvl w:val="0"/>
      <w:numId w:val="1"/>
    </w:numPr>
  </w:pPr>
  <w:r>
    <w:rPr>
      <w:rFonts w:eastAsia="{font_name}" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
      <w:sz w:val="{font_size}"/>
      <w:szCs w:val="{font_size}"/>
    </w:rPr>
    <w:t>{text}</w:t>
  </w:r>
</w:p>""")

        elif elem["type"] == "bold":
            text = escape_xml(elem["text"])
            font_name = config.get("body_font", "宋体")
            font_size = config.get("body_size", 12) * 2
            
            body_parts.append(f"""<w:p>
  <w:pPr>
    <w:spacing w:line="240" w:lineRule="auto"/>
  </w:pPr>
  <w:r>
    <w:rPr>
      <w:rFonts w:eastAsia="{font_name}" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
      <w:b/>
      <w:sz w:val="{font_size}"/>
      <w:szCs w:val="{font_size}"/>
    </w:rPr>
    <w:t>{text}</w:t>
  </w:r>
</w:p>""")

        elif elem["type"] == "quote":
            text = escape_xml(elem["text"])
            font_name = config.get("blockquote_font", "楷体")
            font_size = config.get("blockquote_size", 12) * 2
            line_spacing = int(config.get("line_spacing", 1.25) * 240)
            
            body_parts.append(f"""<w:p>
  <w:pPr>
    <w:spacing w:line="{line_spacing}" w:lineRule="auto"/>
    <w:ind w:left="720"/>
    <w:jc w:val="both"/>
  </w:pPr>
  <w:r>
    <w:rPr>
      <w:rFonts w:eastAsia="{font_name}" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
      <w:sz w:val="{font_size}"/>
      <w:szCs w:val="{font_size}"/>
    </w:rPr>
    <w:t>{text}</w:t>
  </w:r>
</w:p>""")

        elif elem["type"] == "table":
            header = elem["header"]
            rows = elem["rows"]
            font_name = config.get("table_font", "宋体")
            font_size = config.get("table_size", 12) * 2
            num_cols = len(header)
            
            table_parts = []
            table_parts.append(f"<w:tbl><w:tblPr><w:tblStyle w:val='TableGrid'/></w:tblPr><w:tblGrid>")
            for _ in range(num_cols):
                table_parts.append("<w:gridCol w:w='4000'/>")
            table_parts.append("</w:tblGrid>")
            
            table_parts.append("<w:tr>")
            for cell_text in header:
                escaped_text = escape_xml(cell_text)
                table_parts.append(f"""<w:tc>
  <w:tcPr><w:tcW w:w='4000' w:type='dxa'/></w:tcPr>
  <w:p>
    <w:pPr><w:jc w:val='center'/></w:pPr>
    <w:r>
      <w:rPr>
        <w:rFonts w:eastAsia="{font_name}" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
        <w:b/>
        <w:sz w:val="{font_size}"/>
        <w:szCs w:val="{font_size}"/>
      </w:rPr>
      <w:t>{escaped_text}</w:t>
    </w:r>
  </w:p>
</w:tc>""")
            table_parts.append("</w:tr>")
            
            for row_data in rows:
                table_parts.append("<w:tr>")
                for idx, cell_text in enumerate(row_data):
                    if idx >= num_cols:
                        break
                    escaped_text = escape_xml(cell_text)
                    table_parts.append(f"""<w:tc>
  <w:tcPr><w:tcW w:w='4000' w:type='dxa'/></w:tcPr>
  <w:p>
    <w:pPr><w:jc w:val='left'/></w:pPr>
    <w:r>
      <w:rPr>
        <w:rFonts w:eastAsia="{font_name}" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>
        <w:sz w:val="{font_size}"/>
        <w:szCs w:val="{font_size}"/>
      </w:rPr>
      <w:t>{escaped_text}</w:t>
    </w:r>
  </w:p>
</w:tc>""")
                table_parts.append("</w:tr>")
            
            table_parts.append("</w:tbl>")
            body_parts.append("\n".join(table_parts))

    body_content = "\n".join(body_parts)
    
    margin_top = int(config.get("margin_top", 2.54) * 28.35 * 100)
    margin_bottom = int(config.get("margin_bottom", 2.54) * 28.35 * 100)
    margin_left = int(config.get("margin_left", 3.17) * 28.35 * 100)
    margin_right = int(config.get("margin_right", 3.17) * 28.35 * 100)

    document_xml = f"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    {body_content}
    <w:sectPr>
      <w:pgSz w:w="11906" w:h="16838"/>
      <w:pgMar w:top="{margin_top}" w:right="{margin_right}" w:bottom="{margin_bottom}" w:left="{margin_left}" w:header="708" w:footer="708" w:gutter="0"/>
      <w:cols w:space="708"/>
      <w:docGrid w:linePitch="360"/>
    </w:sectPr>
  </w:body>
</w:document>"""

    return document_xml


def generate_docx_zip(output_path, document_xml, config):
    """Generate a valid .docx file using zipfile."""
    content_types_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""

    rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>"""

    doc_rels_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>"""

    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types_xml.encode("utf-8"))
        zf.writestr("_rels/.rels", rels_xml.encode("utf-8"))
        zf.writestr("word/_rels/document.xml.rels", doc_rels_xml.encode("utf-8"))
        zf.writestr("word/document.xml", document_xml.encode("utf-8"))

    print(f"Word document (fallback mode) saved to: {output_path}")


def generate_word_with_python_docx(md_content, config):
    """Generate Word document using python-docx library."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn

    def set_cell_font(cell, font_name, font_size, bold=False):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.font.bold = bold
                run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def set_run_font(run, font_name, font_size, bold=False, color=None):
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def add_heading_styled(doc, text, level, config):
        heading = doc.add_heading(level=level)
        run = heading.add_run(text)
        if level == 0:
            font_name = config.get("title_font", "黑体")
            font_size = config.get("title_size", 15)
        elif level == 1:
            font_name = config.get("heading1_font", "黑体")
            font_size = config.get("heading1_size", 15)
        elif level == 2:
            font_name = config.get("heading2_font", "黑体")
            font_size = config.get("heading2_size", 14)
        elif level == 3:
            font_name = config.get("heading3_font", "黑体")
            font_size = config.get("heading3_size", 12)
        else:
            font_name = config.get("heading4_font", "宋体")
            font_size = config.get("heading4_size", 12)
        set_run_font(run, font_name, font_size, bold=True)

    def add_body_paragraph(doc, text, config):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para_format = para.paragraph_format
        para_format.line_spacing = config.get("line_spacing", 1.25)
        para_format.first_line_indent = Cm(config.get("first_line_indent", 0.74))
        run = para.add_run(text)
        set_run_font(run, config.get("body_font", "宋体"), config.get("body_size", 12))

    def add_table_from_md(doc, header_row, data_rows, config):
        num_cols = len(header_row)
        table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
        table.style = "Table Grid"

        for i, cell_text in enumerate(header_row):
            cell = table.rows[0].cells[i]
            cell.text = cell_text.strip()
            for para in cell.paragraphs:
                para_format = para.paragraph_format
                para_format.line_spacing = config.get("line_spacing", 1.25)
                for run in para.runs:
                    run.font.name = config.get("table_font", "宋体")
                    run.font.size = Pt(config.get("table_size", 12))
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(50, 50, 50)
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), config.get("table_font", "宋体"))

        for row_idx, row_data in enumerate(data_rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx + 1].cells[col_idx]
                    cell.text = cell_text.strip()
                    for para in cell.paragraphs:
                        para_format = para.paragraph_format
                        para_format.line_spacing = config.get("line_spacing", 1.25)
                        for run in para.runs:
                            run.font.name = config.get("table_font", "宋体")
                            run.font.size = Pt(config.get("table_size", 12))
                            run.font.bold = False
                            run._element.rPr.rFonts.set(qn("w:eastAsia"), config.get("table_font", "宋体"))

    def set_page_margins(section, config):
        section.top_margin = Cm(config.get("margin_top", 2.54))
        section.bottom_margin = Cm(config.get("margin_bottom", 2.54))
        section.left_margin = Cm(config.get("margin_left", 3.17))
        section.right_margin = Cm(config.get("margin_right", 3.17))

    doc = Document()

    for section in doc.sections:
        set_page_margins(section, config)

    style = doc.styles["Normal"]
    style.font.name = config.get("body_font", "宋体")
    style.font.size = Pt(config.get("body_size", 12))
    style._element.rPr.rFonts.set(qn("w:eastAsia"), config.get("body_font", "宋体"))

    lines = md_content.split("\n")
    i = 0
    in_table = False
    table_header = []
    table_rows = []

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            if in_table and table_header:
                add_table_from_md(doc, table_header, table_rows, config)
                table_header = []
                table_rows = []
                in_table = False
            i += 1
            continue

        if line.startswith("|") and "|" in line[1:]:
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if not in_table:
                table_header = cells
                in_table = True
            elif all(c.replace("-", "").replace(":", "").strip() == "" for c in cells):
                pass
            else:
                table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table and table_header:
                add_table_from_md(doc, table_header, table_rows, config)
                table_header = []
                table_rows = []
                in_table = False

        if line.startswith("# "):
            add_heading_styled(doc, line[2:], 0, config)
        elif line.startswith("## "):
            add_heading_styled(doc, line[3:], 1, config)
        elif line.startswith("### "):
            add_heading_styled(doc, line[4:], 2, config)
        elif line.startswith("#### "):
            add_heading_styled(doc, line[5:], 3, config)
        elif line.startswith("- ") or line.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(line[2:])
            set_run_font(run, config.get("body_font", "宋体"), config.get("body_size", 12))
        elif line.startswith("**") and line.endswith("**"):
            para = doc.add_paragraph()
            run = para.add_run(line.strip("*"))
            set_run_font(run, config.get("body_font", "宋体"), config.get("body_size", 12), bold=True)
        elif line.startswith("> "):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(1)
            para_format = para.paragraph_format
            para_format.line_spacing = config.get("line_spacing", 1.25)
            para_format.first_line_indent = Cm(config.get("first_line_indent", 0.74))
            run = para.add_run(line[2:])
            set_run_font(run, config.get("blockquote_font", "楷体"), config.get("blockquote_size", 12))
        else:
            add_body_paragraph(doc, line, config)

        i += 1

    if in_table and table_header:
        add_table_from_md(doc, table_header, table_rows, config)

    return doc


def get_next_report_number():
    from datetime import datetime
    today = datetime.now().strftime('%Y-%m-%d')
    output_dir = Path("output")
    if not output_dir.exists():
        return 1
    
    pattern = f"检索报告-{today}-*.docx"
    existing_files = list(output_dir.glob(pattern))
    
    if not existing_files:
        pattern_no_num = f"检索报告-{today}.docx"
        if (output_dir / pattern_no_num).exists():
            return 2
        return 1
    
    max_num = 1
    for f in existing_files:
        try:
            num_str = f.stem.split("-")[-1]
            num = int(num_str)
            if num > max_num:
                max_num = num
        except ValueError:
            pass
    
    return max_num + 1


def main():
    from datetime import datetime
    
    parser = argparse.ArgumentParser(description="Generate Word document from markdown report")
    parser.add_argument("--input", "-i", required=True, help="Input markdown file path")
    parser.add_argument("--output", "-o", help="Output Word file path")
    parser.add_argument("--config", "-c", help="Config JSON file path (optional)")
    args = parser.parse_args()
    
    now = datetime.now()
    today = now.strftime('%Y-%m-%d')
    time_str = now.strftime('%H%M')
    if args.output:
        output_path = Path(args.output)
    else:
        num = get_next_report_number()
        if num == 1:
            output_path = Path(f"output/检索报告-{today}-{time_str}.docx")
        else:
            output_path = Path(f"output/检索报告-{today}-{num}-{time_str}.docx")

    config = DEFAULT_CONFIG.copy()
    if args.config:
        with open(args.config, "r", encoding="utf-8") as f:
            user_config = json.load(f)
            if "word_format" in user_config:
                config.update(user_config["word_format"])
            else:
                config.update(user_config)

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Input file not found: {args.input}")
        sys.exit(1)

    with open(input_path, "r", encoding="utf-8") as f:
        md_content = f.read()

    output_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        from docx import Document
        print("Using python-docx library...")
        doc = generate_word_with_python_docx(md_content, config)
        doc.save(str(output_path))
        print(f"Word document saved to: {output_path}")
    except ImportError:
        print("python-docx not available, using fallback zipfile mode...")
        elements = parse_markdown_to_elements(md_content)
        document_xml = generate_ooxml_document(elements, config)
        generate_docx_zip(str(output_path), document_xml, config)


if __name__ == "__main__":
    main()